import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.run import Run
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table

_CONTROL_CHARS_RE = re.compile(
    "[" + "".join(chr(c) for c in range(32) if c not in (9, 10, 13)) + "]"
)

_UNICODE_NORMALIZATION_MAP = {
    "\u2010": "-",
    "\u2011": "-",
    "\u2012": "-",
    "\u2013": "-",
    "\u2014": "-",
    "\u2212": "-",
    "\u2022": " ",
    "\u00b7": " ",
    "\u25a0": " ",
    "\u25aa": " ",
    "\u00a0": " ",
}


def _clean_text(value: object | None) -> str:
    """
    Normalize arbitrary values to a safe string and remove control characters
    that are invalid in XML/PDF backends.
    """
    if value is None:
        return ""
    # If we get a list/tuple/set (e.g. model returns a list of lines), join them.
    if isinstance(value, list | tuple | set):
        value = ", ".join(str(v) for v in value if v is not None)
    elif not isinstance(value, str):
        value = str(value)

    cleaned = _CONTROL_CHARS_RE.sub("", value)
    for src, replacement in _UNICODE_NORMALIZATION_MAP.items():
        cleaned = cleaned.replace(src, replacement)
    return cleaned


def _contact_link_url(key: str, value: str) -> str | None:
    """
    Return the URL to use for a contact link, or None if it should be plain text.
    Handles cases where only display text was stored (e.g. "LinkedIn", "GitHub")
    instead of the actual URL from the original CV hyperlink.
    """
    if not value or key not in ("linkedin", "website"):
        return None
    raw = value.strip()
    if not raw:
        return None
    lower = raw.lower()
    if raw.startswith("http://") or raw.startswith("https://"):
        return raw
    if key == "linkedin" and lower in ("linkedin", "linkedin profile", "linkedin url"):
        return "https://linkedin.com"
    if lower in ("github", "github profile", "github url"):
        return "https://github.com"
    if key == "website" and "." in raw and " " not in raw:
        return raw if "://" in raw else "https://" + raw
    return None


def _add_hyperlink_run(
    paragraph, text: str, url: str, font_name: str, font_size_pt: int, color_rgb: tuple
):
    """Add a run that is a clickable hyperlink, with the same styling as contact text."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = Run(OxmlElement("w:r"), paragraph)
    new_run.text = text
    new_run.font.name = font_name
    new_run.font.size = Pt(font_size_pt)
    new_run.font.color.rgb = RGBColor(*color_rgb)
    new_run.font.underline = False
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)


def _add_heading(
    doc: Document, text: str, level: int = 1, heading_pt: int = 13, space_before: int = 0
):
    heading = doc.add_heading(_clean_text(text), level=level)
    heading.paragraph_format.space_before = Pt(space_before)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        run.font.name = "Calibri"
        run.font.size = Pt(heading_pt)
    return heading


def _add_paragraph(
    doc: Document,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size: int = 11,
    space_after: int = 2,
):
    p = doc.add_paragraph()
    run = p.add_run(_clean_text(text))
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def _add_bullet(doc: Document, text: str, size: int = 11, space_after: int = 1):
    p = doc.add_paragraph(_clean_text(text), style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def _add_separator(doc: Document, space_before: int = 2, space_after: int = 2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)


def build_cv_docx(cv_data: dict, page_limit: int = 1) -> bytes:
    compact = page_limit == 1
    # Compact mode: tighter than 2-page but enough spacing to fill the page and avoid a large bottom gap
    margin_in = 0.4 if compact else 0.6
    margin_side = 0.5 if compact else 0.75
    name_pt = 20 if compact else 24
    contact_pt = 8 if compact else 10
    body_pt = 10 if compact else 11
    meta_pt = 8 if compact else 9
    heading_pt = 11 if compact else 14
    exp_header_pt = body_pt + 1 if compact else body_pt
    edu_header_pt = body_pt + 1 if compact else body_pt
    space_after_body = 3 if compact else 8
    space_after_bullet = 3 if compact else 6
    separator_before = 3 if compact else 8
    separator_after = 3 if compact else 8
    space_before_heading = 5 if compact else 12
    space_after_block = 5 if compact else 8

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(body_pt)
    style.paragraph_format.space_after = Pt(space_after_body)

    for section in doc.sections:
        section.top_margin = Inches(margin_in)
        section.bottom_margin = Inches(margin_in)
        section.left_margin = Inches(margin_side)
        section.right_margin = Inches(margin_side)

    info = cv_data.get("personal_info", {}) or {}
    name = _clean_text(info.get("full_name", ""))
    if name:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(name)
        run.font.name = "Calibri"
        run.font.size = Pt(name_pt)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    contact_keys = ["email", "phone", "location", "linkedin", "website"]
    contact_items = [(k, _clean_text(info.get(k, ""))) for k in contact_keys]
    contact_items = [(k, v) for k, v in contact_items if v]
    if contact_items:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_color = (0x6B, 0x72, 0x80)
        for i, (key, val) in enumerate(contact_items):
            if i > 0:
                run = contact_para.add_run("   |   ")
                run.font.name = "Calibri"
                run.font.size = Pt(contact_pt)
                run.font.color.rgb = RGBColor(*contact_color)
            link_url = _contact_link_url(key, val)
            if link_url:
                _add_hyperlink_run(
                    contact_para, val, link_url, "Calibri", contact_pt, contact_color
                )
            else:
                run = contact_para.add_run(val)
                run.font.name = "Calibri"
                run.font.size = Pt(contact_pt)
                run.font.color.rgb = RGBColor(*contact_color)

    _add_separator(doc, separator_before, separator_after)

    summary = _clean_text(cv_data.get("professional_summary", ""))
    if summary:
        _add_paragraph(doc, summary, size=body_pt, space_after=space_after_body)

    experience = cv_data.get("experience", []) or []
    if experience:
        _add_heading(
            doc,
            "Professional Experience",
            level=2,
            heading_pt=heading_pt,
            space_before=space_before_heading,
        )
        for exp in experience:
            title_line = _clean_text(exp.get("job_title", ""))
            company = _clean_text(exp.get("company", ""))
            location = _clean_text(exp.get("location", ""))
            dates = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}".strip(" -")

            p = doc.add_paragraph()
            run = p.add_run(title_line)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(exp_header_pt)
            if compact:
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            if company:
                run = p.add_run(f"  |  {company}")
                run.font.name = "Calibri"
                run.font.size = Pt(body_pt)
            meta_parts = []
            if location:
                meta_parts.append(location)
            if dates:
                meta_parts.append(dates)
            if meta_parts:
                run = p.add_run("  |  " + "  |  ".join(meta_parts))
                run.font.name = "Calibri"
                run.font.size = Pt(meta_pt)
                run.italic = True
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            p.paragraph_format.space_after = Pt(space_after_body)

            for bullet in exp.get("bullets", []) or []:
                _add_bullet(doc, bullet, size=body_pt, space_after=space_after_bullet)
        # Space after experience block for visual balance
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after_block)

    education = cv_data.get("education", []) or []
    if education:
        _add_heading(
            doc, "Education", level=2, heading_pt=heading_pt, space_before=space_before_heading
        )
        for edu in education:
            degree = _clean_text(edu.get("degree", ""))
            institution = _clean_text(edu.get("institution", ""))
            dates = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}".strip(" -")

            p = doc.add_paragraph()
            run = p.add_run(degree)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(edu_header_pt)
            if compact:
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            if institution:
                run = p.add_run(f"  |  {institution}")
                run.font.name = "Calibri"
                run.font.size = Pt(body_pt)
            if dates:
                run = p.add_run("  |  " + dates)
                run.font.name = "Calibri"
                run.font.size = Pt(meta_pt)
                run.italic = True
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            p.paragraph_format.space_after = Pt(space_after_body)

            details = _clean_text(edu.get("details", ""))
            if details:
                _add_paragraph(doc, details, size=body_pt, space_after=space_after_body)
        # Space after education block for visual balance
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after_block)

    skills = cv_data.get("skills", {}) or {}
    if any(skills.get(k) for k in ["technical", "soft", "languages", "certifications"]):
        _add_heading(
            doc, "Skills", level=2, heading_pt=heading_pt, space_before=space_before_heading
        )
        for category, label in [
            ("technical", "Technical"),
            ("soft", "Soft Skills"),
            ("languages", "Languages"),
            ("certifications", "Certifications"),
        ]:
            raw_items = skills.get(category, [])
            if isinstance(raw_items, str):
                raw_items = [raw_items]
            items = [_clean_text(str(item)) for item in (raw_items or [])]
            if items:
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(body_pt)
                run = p.add_run(", ".join(items))
                run.font.name = "Calibri"
                run.font.size = Pt(body_pt)

    interests_raw = cv_data.get("interests")
    if interests_raw is not None:
        if isinstance(interests_raw, str):
            interests_list = [interests_raw.strip()] if interests_raw.strip() else []
        else:
            interests_list = [
                _clean_text(str(item)).strip()
                for item in (interests_raw or [])
                if _clean_text(str(item)).strip()
            ]
        if interests_list:
            _add_heading(
                doc, "Interests", level=2, heading_pt=heading_pt, space_before=space_before_heading
            )
            _add_paragraph(
                doc, ", ".join(interests_list), size=body_pt, space_after=space_after_body
            )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_cv_pdf(cv_data: dict, page_limit: int = 1) -> bytes:
    """
    Build a clean, printable PDF version of the CV using the same structured data
    that powers the DOCX export.
    """
    compact = page_limit == 1
    margin_in = 0.4 if compact else 0.6
    margin_side = 0.5 if compact else 0.75
    name_font = 20 if compact else 24
    contact_font = 7 if compact else 9
    heading_font = 11 if compact else 14
    body_font = 10 if compact else 11
    meta_font = 7 if compact else 8
    header_font = body_font + 1 if compact else body_font
    body_leading = 12 if compact else 14
    space_after = 3 if compact else 8
    space_before_heading = 5 if compact else 12
    space_after_block = 5 if compact else 8

    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=margin_side * inch,
        rightMargin=margin_side * inch,
        topMargin=margin_in * inch,
        bottomMargin=margin_in * inch,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CvName",
            parent=styles["Heading1"],
            alignment=TA_CENTER,
            fontSize=name_font,
            fontName="Helvetica-Bold",
            spaceAfter=2 if compact else 8,
            textColor=colors.HexColor("#1E293B"),
        )
    )
    styles.add(
        ParagraphStyle(
            name="CvContact",
            parent=styles["Normal"],
            alignment=TA_CENTER,
            fontSize=contact_font,
            textColor=colors.HexColor("#64748B"),
            spaceAfter=3 if compact else 6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CvHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=heading_font,
            textColor=colors.HexColor("#1E293B"),
            spaceBefore=0,
            spaceAfter=3 if compact else 6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CvBody",
            parent=styles["Normal"],
            fontSize=body_font,
            leading=body_leading,
            spaceAfter=space_after,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CvMeta",
            parent=styles["Normal"],
            fontSize=meta_font,
            textColor=colors.HexColor("#64748B"),
            spaceAfter=space_after,
        )
    )

    story: list = []

    info = cv_data.get("personal_info", {}) or {}
    name = _clean_text(info.get("full_name", ""))
    if name:
        story.append(Paragraph(name, styles["CvName"]))

    contact_keys = ["email", "phone", "location", "linkedin", "website"]
    contact_segments = []
    for key in contact_keys:
        val = _clean_text(info.get(key, ""))
        if not val:
            continue
        link_url = _contact_link_url(key, val)
        escaped = val.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if link_url:
            contact_segments.append(
                '<a href="{}" color="#64748B">{}</a>'.format(
                    link_url.replace("&", "&amp;"), escaped
                )
            )
        else:
            contact_segments.append(escaped)
    if contact_segments:
        story.append(Paragraph("   |   ".join(contact_segments), styles["CvContact"]))

    # Horizontal rule under header
    content_width = (A4[0] / 72.0 - 2 * margin_side) * inch
    rule_table = Table([[""]], colWidths=[content_width], rowHeights=[1])
    rule_table.setStyle([("LINEBELOW", (0, 0), (-1, 0), 0.5, colors.HexColor("#94A3B8"))])
    story.append(Spacer(1, 2))
    story.append(rule_table)
    story.append(Spacer(1, 5 if compact else 12))

    summary = _clean_text(cv_data.get("professional_summary", ""))
    if summary:
        story.append(Paragraph(summary, styles["CvBody"]))

    def _pdf_esc(s: str) -> str:
        return s.replace("&", "&amp;").replace("<", "&lt;")

    experience = cv_data.get("experience", []) or []
    if experience:
        story.append(Spacer(1, space_before_heading))
        story.append(Paragraph("Professional Experience", styles["CvHeading"]))
        for exp in experience:
            title_line = _clean_text(exp.get("job_title", ""))
            company = _clean_text(exp.get("company", ""))
            location = _clean_text(exp.get("location", ""))
            dates = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}".strip(" -")

            line_parts: list[str] = []
            if title_line:
                title_html = _pdf_esc(title_line)
                # Emphasize job headers a bit more, especially on compact (1-page) CVs
                if compact:
                    title_html = '<b><font size="%d">%s</font></b>' % (header_font, title_html)
                else:
                    title_html = f"<b>{title_html}</b>"
                line_parts.append(title_html)
            if company:
                line_parts.append(_pdf_esc(company))
            meta_parts = [p for p in [location, dates] if p]
            if line_parts:
                # Don't escape line_parts here—they already contain safe markup (tags + escaped text)
                line = "  |  ".join(line_parts)
                if meta_parts:
                    line += (
                        '  |  <i><font color="#64748B" size="%d">' % meta_font
                        + "  |  ".join(_pdf_esc(p) for p in meta_parts)
                        + "</font></i>"
                    )
                story.append(Paragraph(line, styles["CvBody"]))

            bullets = [b for b in (exp.get("bullets") or []) if b]
            if bullets:
                items = [
                    ListItem(
                        Paragraph(_clean_text(text), styles["CvBody"]),
                        leftIndent=10,
                    )
                    for text in bullets
                ]
                story.append(
                    ListFlowable(
                        items,
                        bulletType="bullet",
                        start=None,
                        leftIndent=10,
                    )
                )
        story.append(Spacer(1, space_after_block))

    education = cv_data.get("education", []) or []
    if education:
        story.append(Spacer(1, space_before_heading))
        story.append(Paragraph("Education", styles["CvHeading"]))
        for edu in education:
            degree = _clean_text(edu.get("degree", ""))
            institution = _clean_text(edu.get("institution", ""))
            dates = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}".strip(" -")

            line_parts: list[str] = []
            if degree:
                degree_html = _pdf_esc(degree)
                if compact:
                    degree_html = '<b><font size="%d">%s</font></b>' % (header_font, degree_html)
                else:
                    degree_html = f"<b>{degree_html}</b>"
                line_parts.append(degree_html)
            if institution:
                line_parts.append(_pdf_esc(institution))
            if line_parts:
                # Don't escape line_parts—they already contain safe markup (tags + escaped text)
                line = "  |  ".join(line_parts)
                if dates:
                    line += (
                        '  |  <i><font color="#64748B" size="%d">' % meta_font
                        + _pdf_esc(dates)
                        + "</font></i>"
                    )
                story.append(Paragraph(line, styles["CvBody"]))

            details = _clean_text(edu.get("details", ""))
            if details:
                story.append(
                    Paragraph(details.replace("&", "&amp;").replace("<", "&lt;"), styles["CvBody"])
                )
        story.append(Spacer(1, space_after_block))

    skills = cv_data.get("skills", {}) or {}
    if any(skills.get(k) for k in ["technical", "soft", "languages", "certifications"]):
        story.append(Spacer(1, space_before_heading))
        story.append(Paragraph("Skills", styles["CvHeading"]))
        for category, label in [
            ("technical", "Technical"),
            ("soft", "Soft Skills"),
            ("languages", "Languages"),
            ("certifications", "Certifications"),
        ]:
            raw_items = skills.get(category, [])
            if isinstance(raw_items, str):
                raw_items = [raw_items]
            items = [_clean_text(str(item)) for item in (raw_items or [])]
            if items:
                story.append(Paragraph(f"{label}: " + ", ".join(items), styles["CvBody"]))

    interests_raw = cv_data.get("interests")
    if interests_raw is not None:
        if isinstance(interests_raw, str):
            interests_list = [interests_raw.strip()] if interests_raw.strip() else []
        else:
            interests_list = [
                _clean_text(str(item)).strip()
                for item in (interests_raw or [])
                if _clean_text(str(item)).strip()
            ]
        if interests_list:
            story.append(Spacer(1, space_before_heading))
            story.append(Paragraph("Interests", styles["CvHeading"]))
            story.append(
                Paragraph(", ".join(_pdf_esc(t) for t in interests_list), styles["CvBody"])
            )

    if not story:
        story.append(Paragraph("CV data is not available.", styles["CvBody"]))

    doc.build(story)
    return buf.getvalue()


def build_cover_letter_docx(content: str, personal_info: dict | None = None) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    if personal_info:
        name = _clean_text(personal_info.get("full_name", ""))
        if name:
            p = doc.add_paragraph()
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Calibri"

        contact_parts = []
        for key in ["email", "phone", "location"]:
            val = _clean_text(personal_info.get(key, ""))
            if val:
                contact_parts.append(val)
        if contact_parts:
            p = doc.add_paragraph()
            run = p.add_run("  |  ".join(contact_parts))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            run.font.name = "Calibri"

        doc.add_paragraph()

    blocks = re.split(r"\n\s*\n", content) if "\n\n" in content else content.split("\n")
    for block in blocks:
        block = _clean_text(block).strip()
        if block:
            p = doc.add_paragraph()
            run = p.add_run(block)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_cover_letter_pdf(content: str, personal_info: dict | None = None) -> bytes:
    """
    Build a simple, well-formatted PDF version of the cover letter.
    """
    # Normalize line endings so \r\n and \r don't break paragraph splitting or rendering
    content = re.sub(r"\r\n?|\n", "\n", content) if content else ""

    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ClName",
            parent=styles["Heading2"],
            alignment=TA_LEFT,
            fontSize=14,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ClContact",
            parent=styles["Normal"],
            alignment=TA_LEFT,
            fontSize=9,
            textColor=colors.HexColor("#6B7280"),
            spaceAfter=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ClBody",
            parent=styles["Normal"],
            fontSize=11,
            leading=14,
            spaceAfter=8,
        )
    )

    story: list = []

    if personal_info:
        name = _clean_text(personal_info.get("full_name", ""))
        if name:
            story.append(Paragraph(name, styles["ClName"]))

        contact_parts = []
        for key in ["email", "phone", "location"]:
            val = _clean_text(personal_info.get(key, ""))
            if val:
                contact_parts.append(val)
        if contact_parts:
            story.append(Paragraph("  |  ".join(contact_parts), styles["ClContact"]))

        story.append(Spacer(1, 12))

    blocks = re.split(r"\n\s*\n", content) if "\n\n" in content else content.split("\n")
    added_block = False
    for block in blocks:
        block = _clean_text(block).strip()
        if block:
            # Escape so user content is not interpreted as ReportLab markup
            block = block.replace("&", "&amp;").replace("<", "&lt;")
            # Preserve single line breaks within a block (e.g. "Sincerely,\nName")
            block = block.replace("\n", "<br/>")
            story.append(Paragraph(block, styles["ClBody"]))
            added_block = True

    if not added_block:
        story.append(Paragraph("Cover letter content is not available.", styles["ClBody"]))

    doc.build(story)
    return buf.getvalue()
